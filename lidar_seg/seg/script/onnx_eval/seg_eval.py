import docx
import numpy as np
from terminaltables import AsciiTable
import os
import time
import cv2

classes = ('路面', '人行道', '障碍物', '高空噪声', '支撑杆', '交通牌', '交通锥',
           '车或者人', '轮挡', '未知目标',)

palette = [[255, 0, 0], [255, 127, 80], [0, 255, 127],
    [200, 200, 200], [107, 142, 35], [0, 255, 127],
    [152, 251, 152], [0, 0, 255], [142, 0, 252], [128, 64, 128]]

class Job:
    def __init__(self, iou, precision, gt, pred, path, raw_points) -> None:
        self.iou = iou
        self.precision = precision
        self.gt = gt
        self.pred = pred
        self.path = path
        self.raw_points = raw_points
    
    def __eq__(self, other) -> bool:
        return self.iou == other.iou
    
    def __lt__(self, other) -> bool:
        return self.iou > other.iou

class BevVis(object):
    def __init__(self, ranges=(0, 100, -50, 50, -3, 3), voxel_size=(0.4, 0.4, 0.2)):
        self.ranges = ranges
        self.voxel_size = voxel_size

    def generage_bev(self):
        bev_h = (self.ranges[1] - self.ranges[0]) // self.voxel_size[0] + 1
        bev_w = (self.ranges[3] - self.ranges[2]) // self.voxel_size[1] + 1
        bev = np.ones((int(bev_h), int(bev_w), 3)).astype(np.uint8)
        return bev

    def plot_pts(self, lidar, label, p_bev, color_dict, gt_label=False):
        inds = (lidar[:, 0] > self.ranges[0]) & (lidar[:, 0] < self.ranges[1]) & \
               (lidar[:, 1] > self.ranges[2]) & (lidar[:, 1] < self.ranges[3]) & \
               (lidar[:, 2] > self.ranges[4]) & (lidar[:, 2] < self.ranges[5])

        valid = (label >= 0) & inds
        orin_lidar_num = lidar.shape[0]
        lidar = lidar[valid]
        label = label.flatten()[:orin_lidar_num][valid]
        if gt_label:
            label[label == 255] = 0
        bev_h = (self.ranges[1] - self.ranges[0]) // self.voxel_size[0] + 1
        bev_w = (self.ranges[3] - self.ranges[2]) // self.voxel_size[1] + 1
        v = bev_h - (lidar[:, 0] - self.ranges[0]) / self.voxel_size[0]
        u = bev_w - (lidar[:, 1] - self.ranges[2]) / self.voxel_size[1]
        v = np.clip(v, 0, bev_h - 1)
        u = np.clip(u, 0, bev_w - 1)
        v = v.reshape(-1).astype(np.int)
        u = u.reshape(-1).astype(np.int)
        order = np.argsort(lidar[:, 2])
        p_bev[v[order], u[order], ] = [color_dict[i][::-1] for i in label[order]]
        # if kitti: color_dict[i & 0xFF]
        return p_bev

def fast_hist(pred, label, n):
    k = (label >= 0) & (label < n)
    bin_count = np.bincount(
        n * label[k].astype(int) + pred[k], minlength=n ** 2)
    hist = bin_count[:n ** 2].reshape(n, n)
    # do not infer suspension
    hist[3, :] = 0
    hist[:, 3] = 0
    return hist

def per_class_status(hist):
    tp = np.diag(hist)
    fp = hist.sum(1) - tp
    fn = hist.sum(0) - tp
    tp_sum = tp.sum()
    fp_sum = fp.sum()
    m_precision = tp_sum / (tp_sum + fp_sum)
    m_iou = tp_sum / (tp_sum + fp_sum + fp_sum)
    np.seterr(divide='ignore', invalid='ignore')
    p_iou = tp / (fp + fn + tp)
    p_precision = tp / (tp + fp)
    p_recall = tp / (tp + fn)
    return p_iou, p_precision, p_recall, m_iou, m_precision

def lidar_eval(total_hist, label2cat, ignore_index):
    iou, precision, recall, m_iou, m_precision = per_class_status(total_hist[0])
    iou_l1, precision_l1, recall_l1, miou_l1, mprecision_l1 = per_class_status(total_hist[1])
    iou_l2, precision_l2, recall_l2, miou_l2, mprecision_l2 = per_class_status(total_hist[2])
    iou_l3, precision_l3, recall_l3, miou_l3, mprecision_l3 = per_class_status(total_hist[3])

    header_list = [[iou, precision, recall],
                   [iou_l1, precision_l1, recall_l1],
                   [iou_l2, precision_l2, recall_l2],
                   [iou_l3, precision_l3, recall_l3]]

    m_recall = m_precision

    header = ['Classes', 'iou', 'precision', 'recall']
    t_columns = []
    for i in range(len(label2cat)):
        if i == ignore_index:
            continue
        # ret_dict[label2cat[i]] = float(iou[i])
        for j, dist in enumerate(['', '[0,30]', '[30,50]', '[50,]']):
            t_c = [[label2cat[i] + dist]]
            for h in header_list[j]:
                t_c.append([f'{h[i]:.3f}'])
            t_columns.append(t_c)

    t_columns.append([['mean[0,30]'], [f'{miou_l1:.3f}'], [f'{mprecision_l1:.3f}'], [f'{mprecision_l1:.3f}']])
    t_columns.append([['mean[30,50]'], [f'{miou_l2:.3f}'], [f'{mprecision_l2:.3f}'], [f'{mprecision_l2:.3f}']])
    t_columns.append([['mean[50,]'], [f'{miou_l3:.3f}'], [f'{mprecision_l3:.3f}'], [f'{mprecision_l3:.3f}']])
    t_columns.append([['mean'], [f'{m_iou:.3f}'], [f'{m_precision:.3f}'], [f'{m_recall:.3f}']])

    table_data = [header]
    table_rows = [list(zip(*t_col)) for t_col in t_columns]
    for row in table_rows:
        table_data += row
    table = AsciiTable(table_data)
    table.inner_footing_row_border = True
    # print_log('\n' + table.table, logger=logger)
    return table, m_iou

def bev(lidar, label, palette, file_name, gt_label, save_dir):
    img_path = os.path.join(os.path.realpath(save_dir), 'hardcase_image', f'{file_name}.jpg')
    os.makedirs(os.path.dirname(img_path), exist_ok=True)
    if not os.path.exists(img_path):
        print(img_path)
    lidar_vis = BevVis()
    temp_bev = lidar_vis.plot_pts(lidar, label, lidar_vis.generage_bev(), palette, gt_label=gt_label)
    cv2.imwrite(img_path, temp_bev)
    return img_path

def show_worst_instance(raw_cloud, pred, gt, pcd_path, palette, root_dir):
    file_name = pcd_path.split('/')[-1][:-4]
    pred[pred == gt] = -1
    pred_img = bev(raw_cloud, pred, palette, file_name + '_pred', gt_label=False, save_dir=root_dir)
    gt_img = bev(raw_cloud, gt, palette, file_name + '_gt', gt_label=True, save_dir=root_dir)
    return pred_img, gt_img

def calc_hist(result, indices, ignore_index, hists, status):
    gt_label = result['gt_label']
    seg_pred = result['predict_labels']
    gt_seg = gt_label.flatten()
    pred_seg = seg_pred.flatten()
    # filter out ignored points
    pred_seg[gt_seg == ignore_index] = -1
    gt_seg[gt_seg == ignore_index] = -1
    level_1 = indices[0].flatten()
    level_2 = indices[1].flatten()
    level_3 = indices[2].flatten()
    # calculate one instance result
    hist_per_frame = fast_hist(pred_seg, gt_seg, ignore_index)
    hist_per_frame_l1 = fast_hist(pred_seg[level_1], gt_seg[level_1], ignore_index)
    hist_per_frame_l2 = fast_hist(pred_seg[level_2], gt_seg[level_2], ignore_index)
    hist_per_frame_l3 = fast_hist(pred_seg[level_3], gt_seg[level_3], ignore_index)
    status.put(Job(*per_class_status(hist_per_frame)[3:], gt_seg, pred_seg,
    result['path'], result['raw_points']))
    if (len(status.queue) > 10):
        status.get()
        status.task_done()
    hists[0] += hist_per_frame
    hists[1] += hist_per_frame_l1
    hists[2] += hist_per_frame_l2
    hists[3] += hist_per_frame_l3

def evaluate(worst_dict, total_hist, root_dir, ignore, doc_path):
    label2cat = {
        i: cat_name
        for i, cat_name in enumerate(classes)
    }
    table, m_iou = lidar_eval(total_hist, label2cat, ignore)

    os.makedirs(root_dir, exist_ok=True)
    # write to doc
    doc2 = docx.Document()
    time_str = time.strftime('%Y%m%d')
    doc2.add_heading(f'lidar seg evaluate, time {time_str}', 0)
    doc2.add_paragraph('Evaluate Table: ')
    # doc2.add_paragraph(table.table)
    doc_table = doc2.add_table(rows=len(table.table_data), cols=len(table.table_data[0]), style='Table Grid')
    red = docx.shared.RGBColor(0xff, 0, 0)
    green = docx.shared.RGBColor(0, 0xa5, 0)
    for i in range(len(table.table_data)):
        doc_table.cell(i, 0).text = table.table_data[i][0]
    for j in range(1, len(table.table_data[0])):
        doc_table.cell(0, j).text = table.table_data[0][j]
    for i in range(1, len(table.table_data)):
        for j in range(1, len(table.table_data[0])):
            cell = doc_table.cell(i, j)
            cell.paragraphs[0].add_run(table.table_data[i][j])
    
    doc2.add_paragraph('\nConfusion matrix: ')
    str_doc = ['total:', '0-30:', '30-50:', '50+:']
    class_nums = len(total_hist[0]) + 1
    for i in range(len(total_hist)):
        doc2.add_paragraph(str_doc[i])
        doc_table = doc2.add_table(rows=class_nums, cols=class_nums, style='Table Grid')
        for k in range(1, class_nums):
            doc_table.cell(k, 0).text = str(k-1)
        for k in range(1, class_nums):
            doc_table.cell(0, k).text = str(k-1)
        for k, matrix in enumerate(total_hist[i]):
            total = sum(matrix)
            if total == 0:
                for l, t in enumerate(matrix):
                    doc_table.cell(k+1, l+1).text = str(t)
            else:
                for l, t in enumerate(matrix):
                    doc_table.cell(k+1, l+1).text = str(round(t / total * 100))     
    doc2.add_paragraph('\nHard Case Info: ')
    while len(worst_dict.queue) > 0:
        task = worst_dict.get()
        doc2.add_paragraph('For Metric: iou', style='List Bullet')
        metric_value = task.iou
        raw_cloud = task.raw_points
        pred = task.pred
        gt = task.gt
        valid = raw_cloud[3, :] > 0.1
        raw_cloud = raw_cloud[:3, valid].T
        valid = valid.reshape(-1)
        pcd_path = task.path
        img_pred, img_gt = show_worst_instance(
            raw_cloud, pred[valid], gt[valid], pcd_path, palette, root_dir)
        string_1 = 'The {} worst frame is: {}'.format(len(worst_dict.queue) + 1, pcd_path)
        doc2.add_paragraph(string_1)
        string_2 = 'frame iou: {:.4f} ; while total validations iou: {:.4f}' \
            .format(metric_value, m_iou)
        doc2.add_paragraph(string_2)
        doc2.add_paragraph('Predict image:')
        doc2.add_picture(img_pred)
        doc2.add_paragraph('Ground truth image:')
        doc2.add_picture(img_gt)
        worst_dict.task_done()
    doc2.save(f'{root_dir}/eval_table.docx')
