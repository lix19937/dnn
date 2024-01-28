from .lidar_eval import lidar_eval, show_worst_instance
import numpy as np
import docx
import os
import time

import os
import matplotlib.pyplot as plt


def plot_loss(root_dir, save_dir):
    file_path = os.path.join(root_dir, 'exp', 'log.txt')
    if os.path.exists(file_path):
        log_file = open(file_path, 'r')
        data = log_file.readlines()
        log_file.close()
        data = [_.split('|') for _ in data]
        train_loss = [float(_[1][4:].strip()) for _ in data]
        val_loss = [float(_[len(data[0]) -1][5:].strip()) for _ in data if len(_) > len(data[0])]

        x = [_ for _ in range(1, len(train_loss) + 1)]
        x1 = [_ * 10 for _ in range(1, len(val_loss) + 1)]

        plt.plot(x, train_loss, 'g-', label= 'train')          # 设定多组数据；线条颜色和样式
        plt.plot(x1, val_loss, 'r-', label= 'val')
        plt.legend(loc='upper left', bbox_to_anchor=(1,1))  # 像前面这样写是放在图外 # 表示多维数组中label在图表中位置
        plt.tight_layout(pad=5)                            # 设定空白的留白大小---将pad改为h_pad或w_pad可以分别设定高度或宽度的留白

        plt.axis([0, len(train_loss), 0, 1])                                    # 产生x刻度0~8、y刻度0~70
        plt.grid()                                                 # 产生网格
        plt.title('Loss Chart', fontsize=24)                       # 图标标题
        plt.xlabel('Epoch', fontsize=16)                           # x轴标题 --- fontsize 设置标题字体大小
        plt.ylabel('Loss', fontsize=16)                          # y轴标题 ---- xticks()设定刻度
        plt.tick_params(axis='both', color='red', labelsize=12)
        # axis中的both是应用到x轴和y轴，如果both替换成x则是仅替换x轴，y就是y轴；color是设定刻度的颜色；labelsize是设定刻度大小
        save_path = os.path.join(save_dir, 'loss.jpg')
        plt.savefig(save_path, bbox_inches='tight')
        return save_path

def compare_table(table_data, lines):
    old_data = np.array([[float(_) for _ in line[1:]] for line in lines[1:]])
    new_data = np.array([[float(_) for _ in line[1:]] for line in table_data[1:]])
    diff = new_data - old_data
    return diff

def evaluate(opt,
             worst_dict,
             total_hist,
             classes,
             palette,
             root_dir):
        """Evaluation in semantic segmentation protocol.

        Args:
            results (list[dict]): List of results.
            metric (str | list[str]): Metrics to be evaluated.
            logger (logging.Logger | None | str): Logger used for printing
                related information during evaluation. Defaults to None.

        Returns:
            dict: Evaluation results.
        """

        label2cat = {
            i: cat_name
            for i, cat_name in enumerate(classes)
        }
        table, m_iou = lidar_eval(total_hist, label2cat, opt.ignore_index)

        lines = [['NaN' for col in range(len(table.table_data[0]))] for row in range(len(table.table_data))]
        save_dir = os.path.dirname(opt.load_model)
        try:
            doc1 = docx.Document(f'{save_dir}/eval_table.docx')
            doc_table = doc1.tables[0]
            for i in range(len(lines)):
                for j in range(len(lines[0])):
                    lines[i][j] = doc_table.cell(i, j).text
        except:
            pass
        
        # with open(f'{root_dir}/eval_table.txt', 'w') as f:
        #     f.writelines(table.table)
        
        
        os.makedirs(root_dir, exist_ok=True)
        # write to doc
        diff = compare_table(table.table_data, lines)
        doc2 = docx.Document()
        time_str = time.strftime('%Y%m%d')
        doc2.add_heading(f'lidar seg evaluate, time {time_str}', 0)
        doc2.add_paragraph('Evaluate Table: ')
        # doc2.add_paragraph(table.table)
        doc_table = doc2.add_table(rows=len(table.table_data), cols=len(table.table_data[0]), style='Table Grid')
        red = docx.shared.RGBColor(0xff, 0, 0)
        green = docx.shared.RGBColor(0, 0xa5, 0)
        for i in range(len(table.table_data)):
            doc_table.cell(i, 0).text = table.table_data[i][0]
        for j in range(1, len(table.table_data[0])):
            doc_table.cell(0, j).text = table.table_data[0][j]
        for i in range(1, len(table.table_data)):
            for j in range(1, len(table.table_data[0])):
                cell = doc_table.cell(i, j)
                run = cell.paragraphs[0].add_run(table.table_data[i][j])
                if diff[i-1][j-1] == 0 or np.isnan(diff[i-1][j-1]):
                    continue
                run.font.color.rgb = green if diff[i - 1][j - 1] > 0 else red

        doc2.add_paragraph('Diff Table: ')
        # doc2.add_paragraph(table.table)
        doc_table = doc2.add_table(rows=len(table.table_data), cols=len(table.table_data[0]), style='Table Grid')
        for i in range(len(table.table_data)):
            doc_table.cell(i, 0).text = table.table_data[i][0]
        for j in range(1, len(table.table_data[0])):
            doc_table.cell(0, j).text = table.table_data[0][j]
        for i in range(1, len(table.table_data)):
            for j in range(1, len(table.table_data[0])):
                cell = doc_table.cell(i, j)
                run = cell.paragraphs[0].add_run(f'{diff[i-1][j-1]:.3f}')
                if diff[i-1][j-1] == 0 or np.isnan(diff[i-1][j-1]):
                    continue
                run.font.color.rgb = green if diff[i - 1][j - 1] > 0 else red

        doc2.add_paragraph('\nConfusion matrix: ')
        str_doc = ['total:', '0-30:', '30-50:', '50+:']
        class_nums = len(total_hist[0]) + 1
        for i in range(len(total_hist)):
            doc2.add_paragraph(str_doc[i])
            doc_table = doc2.add_table(rows=class_nums, cols=class_nums, style='Table Grid')
            for k in range(1, class_nums):
                doc_table.cell(k, 0).text = str(k-1)
            for k in range(1, class_nums):
                doc_table.cell(0, k).text = str(k-1)
            for k, matrix in enumerate(total_hist[i]):
                total = sum(matrix)
                if total == 0:
                    for l, t in enumerate(matrix):
                        doc_table.cell(k+1, l+1).text = str(t)
                else:
                    for l, t in enumerate(matrix):
                        doc_table.cell(k+1, l+1).text = str(round(t / total * 100))     

        doc2.add_paragraph('\nLoss Chart: ')
        plot = plot_loss(opt.root_dir, root_dir)
        # doc2.add_picture(plot)
        doc2.add_paragraph('\nHard Case Info: ')
        # for key in worst_dict.keys():
        while len(worst_dict.queue) > 0:
            task = worst_dict.get()
            doc2.add_paragraph('For Metric: iou', style='List Bullet')
            metric_value = task.iou
            raw_cloud = task.raw_points
            pred = task.pred
            gt = task.gt
            valid = raw_cloud[3, :] > 0.1
            raw_cloud = raw_cloud[:3, valid].T
            valid = valid.reshape(-1)
            pcd_path = task.path
            img_pred, img_gt = show_worst_instance(
                raw_cloud, pred[valid], gt[valid], pcd_path, palette, root_dir)
            string_1 = 'The {} worst frame is: {}'.format(len(worst_dict.queue) + 1, pcd_path)
            doc2.add_paragraph(string_1)
            string_2 = 'frame iou: {:.4f} ; while total validations iou: {:.4f}' \
                .format(metric_value, m_iou)
            doc2.add_paragraph(string_2)
            doc2.add_paragraph('Predict image:')
            doc2.add_picture(img_pred, width=docx.shared.Inches(3.5))
            doc2.add_paragraph('Ground truth image:')
            doc2.add_picture(img_gt, width=docx.shared.Inches(3.5))
            worst_dict.task_done()

        doc2.save(f'{root_dir}/eval_table.docx')
        # return ret_dict
